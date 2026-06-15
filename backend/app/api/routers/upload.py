from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Request
from typing import Dict
import os
from datetime import datetime
import logging
from pathlib import Path
from app.rag.documents.document_rag import DocumentRAG
from app.tools.documents.summarize_document import summarize_document_tool as summarize_document
from fastapi.responses import FileResponse
from app.core.security import get_current_user
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from PIL import Image, ImageEnhance, ImageFilter
from app.rag.documents.document_loader import DocumentLoader
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
import shutil

TESSERACT_AVAILABLE = shutil.which("tesseract") is not None
logger = logging.getLogger(__name__)
router = APIRouter(prefix="/upload", tags=["Upload"])

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def preprocess_and_ocr(image) -> str:
    try:
        if not TESSERACT_AVAILABLE:
            return "[OCR disabled: Tesseract not installed]"

        if image.mode != 'RGB':
            image = image.convert('RGB')

        enhancer = ImageEnhance.Contrast(image)
        image = enhancer.enhance(2.0)

        enhancer = ImageEnhance.Sharpness(image)
        image = enhancer.enhance(2.0)

        image = image.filter(ImageFilter.MedianFilter())

        config = r'--oem 3 --psm 6 -l vie+eng'

        try:
            return pytesseract.image_to_string(image, config=config).strip()
        except Exception as e:
            logger.warning(f"OCR failed: {e}")
            return "[OCR failed but file uploaded successfully]"

    except Exception as e:
        logger.error(f"OCR preprocessing error: {e}")
        return "[OCR error]"


def extract_text_from_file(file: UploadFile) -> str:
    """Extract text từ nhiều định dạng file"""
    filename = file.filename.lower() if file.filename else ""
    content = ""

    try:
        if filename.endswith(".pdf"):
            from PyPDF2 import PdfReader
            reader = PdfReader(file.file)
            for page in reader.pages:
                text = page.extract_text() or ""
                content += text + "\n"

        elif filename.endswith(".docx"):
            from docx import Document as DocxDocument
            doc = DocxDocument(file.file)
            for para in doc.paragraphs:
                content += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    content += row_text + "\n"

        elif filename.endswith((".xlsx", ".xls")):
            import openpyxl
            wb = openpyxl.load_workbook(file.file, data_only=True)
            for sheet in wb.worksheets:
                content += f"\n===== SHEET: {sheet.title} =====\n"
                for row in sheet.iter_rows(values_only=True):
                    values = [str(cell) if cell is not None else "" for cell in row]
                    content += " | ".join(values) + "\n"

        elif filename.endswith((".png", ".jpg", ".jpeg", ".webp")):
            image = Image.open(file.file)
            try:
                content = preprocess_and_ocr(image)
            except Exception as e:
                logger.warning(f"OCR skipped: {e}")
                content = "[Image uploaded but OCR unavailable]"

            if len(content.strip()) < 30:
                content = "[OCR: Nội dung quá ít hoặc ảnh mờ.]"

        else:
            raise HTTPException(
                status_code=400,
                detail="Định dạng file không được hỗ trợ."
            )

        content = " ".join(content.split())
        return content.strip()

    except Exception as e:
        logger.error(f"Extraction error {filename}: {e}")
        raise HTTPException(status_code=400, detail=f"Lỗi đọc file {filename}: {str(e)}")


@router.post("/")
async def upload_file(
    request: Request,                                 # <-- ĐÃ THÊM: Tiếp nhận Request để lắng nghe Frontend hủy mạch
    file: UploadFile = File(...),
    current_user: str = Depends(get_current_user)
) -> Dict:

    if not file.filename:
        raise HTTPException(status_code=400, detail="Không có file")

    file_path = None
    try:
        # 🛑 ĐIỂM DỪNG 1: Kiểm tra trước khi thực hiện trích xuất và bóc tách cấu trúc tài liệu nặng
        if await request.is_disconnected():
            logger.warning(f"🛑 [UPLOAD CANCELLED] Người dùng hủy kết nối trước khi phân tích file: {file.filename}")
            return {"status": "cancelled", "message": "Yêu cầu upload đã bị hủy bởi người dùng."}

        # === SỬ DỤNG DOCUMENT LOADER MỚI (page-aware) ===
        docs = await DocumentLoader.load_file_to_documents(file, current_user)

        # 🛠️ THAY ĐỔI TẠI ĐÂY: Lưu file bằng tên gốc của người dùng để đồng bộ với Data Preview
        file_path = UPLOAD_DIR / file.filename

        file.file.seek(0)
        with open(file_path, "wb") as f:
            f.write(await file.read())

        # 🛠️ THAY ĐỔI TẠI ĐÂY: Sửa đường dẫn URL để khớp chính xác với router GET bên dưới (/upload/uploads/...)
        file_url = f"/upload/uploads/{file_path.name}"

        # 🛑 ĐIỂM DỪNG 2: Kiểm tra trước khi nén vào Vector Store (Tránh tốn tài nguyên nhúng vector và tạo 172 chunks)
        if await request.is_disconnected():
            logger.warning(f"🛑 [UPLOAD CANCELLED] Người dùng hủy kết nối trước khi lưu Vector Store: {file.filename}")
            if file_path and file_path.exists():
                file_path.unlink()  # Dọn dẹp file vật lý vừa lưu tạm để tránh rác ổ cứng
            return {"status": "cancelled", "message": "Yêu cầu upload đã bị hủy trước khi lưu vào cơ sở dữ liệu."}

        # Index vào vector store
        doc_rag = DocumentRAG()
        doc_rag.add_documents(docs)

        logger.info(f"✅ Uploaded & indexed: {file.filename} → {len(docs)} documents")

        # 🛑 ĐIỂM DỪNG 3: Kiểm tra trước khi gọi tiến trình tóm tắt tự động qua LLM (Tiết kiệm Token Groq)
        if await request.is_disconnected():
            logger.warning(f"🛑 [UPLOAD CANCELLED] Đã index xong nhưng người dùng hủy kết nối trước khi chạy Auto-Summarize: {file.filename}")
            return {
                "status": "success",
                "filename": file.filename,
                "doc_id": docs[0].metadata.get("doc_id") if docs else None,
                "chunks": len(docs),
                "file_url": file_url,
                "summary": "Tiến trình tóm tắt tự động bị hủy bởi người dùng."
            }

        # Tóm tắt nhanh sau khi upload
        summary = "Tài liệu đã được upload thành công."
        try:
            summary_result = summarize_document.invoke({"filename": file.filename})
            if isinstance(summary_result, dict):
                if "data" in summary_result and "answer" in summary_result["data"]:
                    summary = summary_result["data"]["answer"]
                else:
                    summary = summary_result.get("summary") or summary_result.get("answer", "")
        except Exception as sum_err:
            logger.warning(f"Auto-summary tool missed: {sum_err}")

        return {
            "status": "success",
            "filename": file.filename,
            "doc_id": docs[0].metadata.get("doc_id") if docs else None,
            "chunks": len(docs),
            "file_url": file_url,
            "summary": summary
        }

    except Exception as e:
        logger.error(f"Upload error: {e}", exc_info=True)
        # Khôi phục/dọn dẹp file thừa nếu gặp lỗi hệ thống giữa chừng
        if file_path and file_path.exists():
            try:
                file_path.unlink()
            except:
                pass
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/uploads/{filename}")
async def get_uploaded_file(filename: str):
    file_path = UPLOAD_DIR / filename
    if not file_path.exists():
        raise HTTPException(404, "File not found")
    
    # 🛠️ THAY ĐỔI TẠI ĐÂY: Thêm tham số filename vào FileResponse để ép trình duyệt/Data Preview hiển thị đúng tên gốc
    return FileResponse(file_path, filename=filename)