# backend/app/rag/documents/document_loader.py
import os
import logging
from typing import List, Dict, Any
from pathlib import Path
from datetime import datetime

from langchain_core.documents import Document
from fastapi import UploadFile

logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    Document Loader chuyên nghiệp - Hỗ trợ nhiều định dạng file
    Dùng cho Upload và Document Agent
    """

    @staticmethod
    def preprocess_and_ocr(image) -> str:
        """OCR tối ưu cho tiếng Việt"""
        try:
            from PIL import Image, ImageEnhance, ImageFilter
            import pytesseract

            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Tăng chất lượng ảnh
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2.0)
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2.0)
            image = image.filter(ImageFilter.MedianFilter())

            # Cấu hình Tesseract cho tiếng Việt
            custom_config = r'--oem 3 --psm 6 -l vie+eng'
            
            text = pytesseract.image_to_string(
                image, 
                config=custom_config,
                lang='vie+eng'
            )
            return text.strip()

        except Exception as e:
            logger.warning(f"OCR preprocessing failed: {e}")
            # Fallback
            try:
                import pytesseract
                return pytesseract.image_to_string(image, lang='vie+eng').strip()
            except:
                return "[OCR Error: Không thể đọc text từ ảnh]"

    @staticmethod
    async def extract_text_from_file(file: UploadFile) -> str:
        """
        Extract text từ nhiều loại file
        """
        filename = file.filename.lower()
        content = ""

        try:
            if filename.endswith(".pdf"):
                from PyPDF2 import PdfReader
                reader = PdfReader(file.file)
                for page in reader.pages:
                    text = page.extract_text() or ""
                    content += text + "\n"

            elif filename.endswith(".docx"):
                from docx import Document
                doc = Document(file.file)
                for para in doc.paragraphs:
                    content += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        content += row_text + "\n"

            elif filename.endswith((".xlsx", ".xls")):
                import openpyxl
                wb = openpyxl.load_workbook(file.file, data_only=True)
                for sheet in wb.worksheets:
                    content += f"\n===== SHEET: {sheet.title} =====\n"
                    for row in sheet.iter_rows(values_only=True):
                        values = [str(cell) if cell is not None else "" for cell in row]
                        content += " | ".join(values) + "\n"

            elif filename.endswith((".png", ".jpg", ".jpeg", ".webp")):
                from PIL import Image
                import pytesseract

                tesseract_cmd = os.getenv("TESSERACT_CMD")
                if tesseract_cmd:
                    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

                image = Image.open(file.file)
                content = DocumentLoader.preprocess_and_ocr(image)

                if len(content.strip()) < 30:
                    content = "[OCR: Ảnh mờ hoặc không có chữ. Hãy thử ảnh rõ nét hơn.]"

            else:
                raise ValueError(f"Định dạng file không hỗ trợ: {filename}")

            # Post-processing
            content = " ".join(content.split())  # Giảm khoảng trắng thừa
            content = content.replace("Sir dung", "Sử dụng").replace("chit ma tity", "chất ma túy")

            return content.strip()

        except ImportError as e:
            logger.error(f"Missing library: {e}")
            raise RuntimeError(f"Thiếu thư viện: {str(e)}")
        except Exception as e:
            logger.error(f"Extract text error {filename}: {e}")
            raise RuntimeError(f"Lỗi đọc file {filename}: {str(e)}")

    @staticmethod
    async def _process_pdf_pages(file: UploadFile) -> List[Document]:
        """Xử lý PDF theo từng trang (Page-aware) - Tối ưu cho RAG"""
        documents = []
        try:
            import pdfplumber
            import io

            file.file.seek(0)
            pdf_bytes = file.file.read()
            file.file.seek(0)  # Reset pointer cho các bước sau

            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text() or ""
                    
                    # Xử lý bảng
                    tables = page.extract_tables()
                    table_text = ""
                    if tables:
                        for idx, table in enumerate(tables):
                            if table:
                                table_text += f"\n[BẢNG {idx+1}]\n"
                                table_text += "\n".join([" | ".join(str(cell) if cell is not None else "" for cell in row) for row in table])
                                table_text += "\n[/BẢNG]\n"

                    content = (text + "\n" + table_text).strip()

                    metadata = {
                        "type": "pdf_page",
                        "filename": file.filename,
                        "page_number": page_num,
                        "total_pages": len(pdf.pages),
                        "file_type": "PDF",
                        "source": "user_upload",
                        "upload_time": datetime.now().isoformat(),
                        "content_length": len(content),
                    }

                    documents.append(Document(page_content=content, metadata=metadata))

            logger.info(f"✅ Processed PDF {file.filename} → {len(documents)} pages")
            return documents

        except Exception as e:
            logger.warning(f"pdfplumber failed for {file.filename}, fallback to PyPDF2: {e}")
            # Fallback về cách cũ nếu pdfplumber lỗi
            content = await DocumentLoader.extract_text_from_file(file)
            file.file.seek(0)
            
            metadata = {
                "type": "pdf_full",
                "filename": file.filename,
                "file_type": "PDF",
                "source": "user_upload",
                "upload_time": datetime.now().isoformat(),
                "content_length": len(content),
                "note": "fallback_full_document"
            }
            return [Document(page_content=content, metadata=metadata)]

    @staticmethod
    async def load_file_to_document(
        file: UploadFile, 
        current_user: str = "system"
    ) -> Document:
        """
        Load file và chuyển thành Langchain Document với metadata đầy đủ
        (Giữ hàm cũ để tương thích)
        """
        content = await DocumentLoader.extract_text_from_file(file)
        
        # Reset file pointer để có thể lưu sau
        file.file.seek(0)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        doc_id = f"upload_{timestamp}"

        metadata = {
            "type": "uploaded_file",
            "filename": file.filename,
            "uploaded_by": current_user,
            "upload_time": datetime.now().isoformat(),
            "source": "user_upload",
            "doc_id": doc_id,
            "file_type": file.filename.split('.')[-1].upper(),
            "content_length": len(content),
        }

        return Document(
            page_content=content,
            metadata=metadata
        )

    @staticmethod
    async def load_file_to_documents(
        file: UploadFile, 
        current_user: str = "system"
    ) -> List[Document]:
        """
        Load file và trả về danh sách Document (page-aware)
        Đây là hàm chính nên dùng cho RAG mới
        """
        filename = file.filename.lower()
        
        try:
            if filename.endswith(".pdf"):
                docs = await DocumentLoader._process_pdf_pages(file)
            elif filename.endswith((".png", ".jpg", ".jpeg", ".webp")):
                # Xử lý ảnh thành 1 document
                content = await DocumentLoader.extract_text_from_file(file)
                file.file.seek(0)
                
                metadata = {
                    "type": "image",
                    "filename": file.filename,
                    "uploaded_by": current_user,
                    "upload_time": datetime.now().isoformat(),
                    "source": "user_upload",
                    "file_type": file.filename.split('.')[-1].upper(),
                    "content_length": len(content),
                }
                docs = [Document(page_content=content, metadata=metadata)]
            else:
                # Các loại khác dùng hàm cũ
                doc = await DocumentLoader.load_file_to_document(file, current_user)
                docs = [doc]

            # Thêm metadata chung cho tất cả
            for doc in docs:
                if "uploaded_by" not in doc.metadata:
                    doc.metadata["uploaded_by"] = current_user
                if "doc_id" not in doc.metadata:
                    doc.metadata["doc_id"] = f"upload_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

            return docs

        except Exception as e:
            logger.error(f"Load file error {file.filename}: {e}")
            raise

    @staticmethod
    async def load_multiple_files(files: List[UploadFile], current_user: str = "system") -> List[Document]:
        """Load nhiều file cùng lúc"""
        documents = []
        for file in files:
            try:
                # Sử dụng hàm mới để hỗ trợ page-aware
                docs = await DocumentLoader.load_file_to_documents(file, current_user)
                documents.extend(docs)
            except Exception as e:
                logger.error(f"Failed to load {file.filename}: {e}")
                continue
        return documents 